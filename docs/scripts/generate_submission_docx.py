#!/usr/bin/env python3
"""Generate Blue Machines case study and internal notes DOCX files."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "submission"
DOWNLOADS = Path.home() / "Downloads"


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def add_title(doc: Document, text: str, subtitle: str = "") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.italic = True
        r2.font.size = Pt(12)
    doc.add_paragraph()


def add_h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_p(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def build_case_study() -> Document:
    doc = Document()
    style_doc(doc)
    add_title(
        doc,
        "Blue Machines — AI Product Operations & Solutions",
        "Case Study Submission · Tiger Credit Card Onboarding Voice Agent · Rushil Kaul",
    )

    add_h(doc, "1. Executive Summary")
    add_p(
        doc,
        "Objective: Design how a Tiger Credit Card AI voice agent integrates with internal systems to move "
        "approved customers through onboarding (eKYC → VKYC → in-app activation), focusing on data flow, "
        "prompt logic, and evaluation — not STT/TTS engineering.",
    )
    add_p(
        doc,
        "Deliverables: (1) System interaction design with READ/WRITE/NOTIFY/ESCALATE flows across 10 systems "
        "and 8 journey stages; (2) Full VAPI system prompt with stage routing and objection logic; "
        "(3) LLM-as-judge eval prompt + operational metrics framework; (4) Interactive prototype for panel review.",
    )
    add_bullets(
        doc,
        [
            "GitHub: https://github.com/TheRuKa7/bluemachines_sample",
            "Interactive prototype: artifacts/tiger-prototype (deploy via Netlify)",
            "Full VAPI prompt: docs/VAPI_SYSTEM_PROMPT.md",
            "Eval prompt: docs/EVAL_PROMPT.md",
        ],
    )

    add_h(doc, "2. Assumptions")
    add_bullets(
        doc,
        [
            "Tiger has Card Core, CRM/Journey Orchestrator, eKYC, VKYC, Activation, Notification, Inside Sales, Compliance, and Analytics systems.",
            "Agent is triggered on drop-off (app stall or missed deadline); Inside Sales may also initiate.",
            "Blue Machines is the orchestration layer — reads state, runs conversation, writes disposition, triggers notifications.",
            "Single source of truth for journey stage is CRM/API — agent never infers stage from chat alone.",
            "Assignment scope excludes telephony pipeline, STT orchestration, and integration engineering.",
        ],
    )

    add_h(doc, "3. Task 1 — System Interaction Design")
    add_h(doc, "3.1 Architecture pattern", 2)
    add_p(doc, "Event-driven, stateful workflow. BM Agent sits at the centre and orchestrates reads from source systems, conversation, writes to CRM/Compliance/Analytics, and notifications/escalations.")

    add_h(doc, "3.2 Ten systems", 2)
    add_table(
        doc,
        ["System", "Role", "Primary operations"],
        [
            ["Card Core", "Product truth", "READ: approval, fees, rewards, limit"],
            ["CRM / Journey State", "Customer + stage SoT", "READ/WRITE: stage, disposition, history"],
            ["eKYC", "Electronic KYC", "READ/WRITE: status, failure, retry"],
            ["VKYC", "Video KYC", "READ/WRITE: slots, no-show, completion"],
            ["Activation", "In-app activation", "READ/WRITE: eligibility, virtual card"],
            ["Blue Machines Agent", "Voice orchestration", "READ all · WRITE outcomes"],
            ["Notification", "WhatsApp/SMS/push", "NOTIFY: deep-links, reminders"],
            ["Inside Sales", "Human fallback", "ESCALATE: context bundle"],
            ["Compliance / Audit", "Guardrails", "WRITE: consent, PII access log"],
            ["Analytics / Eval", "Learning loop", "WRITE: metrics, eval scores"],
        ],
    )

    add_h(doc, "3.3 Eight journey stages", 2)
    add_table(
        doc,
        ["Stage", "Agent goal", "Key READ", "Key WRITE/NOTIFY", "Transition event"],
        [
            ["APPROVED", "Start eKYC", "approval_status, consent", "NOTIFY eKYC link · WRITE CRM", "eKYC initiated"],
            ["EKYC_PENDING", "Recover drop-off", "ekyc_status, retry_count", "WRITE eKYC re-trigger", "eKYC complete"],
            ["EKYC_COMPLETE", "Schedule VKYC", "available_vkyc_slots", "WRITE slot · NOTIFY confirm", "VKYC booked"],
            ["VKYC_PENDING", "Reduce no-show", "vkyc_slot_time, no_show", "NOTIFY reminder · ESCALATE if >2", "VKYC complete"],
            ["VKYC_COMPLETE", "Guide activation", "activation_eligibility", "NOTIFY activation link", "Activation done"],
            ["ACTIVATION_PENDING", "Resolve blocker", "activation_failure_reason", "WRITE retry · ESCALATE tech", "Card active"],
            ["ACTIVE", "Close journey", "welcome eligibility", "WRITE journey closed · NOTIFY welcome", "None"],
            ["ESCALATED", "Human handoff", "escalation_reason, history", "ESCALATE Inside Sales", "Human resolves"],
        ],
    )

    add_h(doc, "3.4 Arrow semantics", 2)
    add_bullets(
        doc,
        [
            "READ (blue): Agent loads structured state before/during call",
            "WRITE (green): Agent or platform persists outcome, disposition, audit",
            "NOTIFY (amber): WhatsApp/SMS/push deep-link or reminder",
            "ESCALATE (red): Context bundle to Inside Sales with full history",
        ],
    )

    add_h(doc, "3.5 Failure & compliance (bonus)", 2)
    add_bullets(
        doc,
        [
            "System timeout → no guessing; callback + reason code + retry ticket",
            "Affected systems highlighted per stage in prototype Failure Mode",
            "PII fields flagged; Compliance WRITE on every outbound contact",
            "Approved product copy only — no improvised fee/reward terms",
            "Escalation thresholds: eKYC retry>3, VKYC no_show>2, repeated objection×2",
        ],
    )

    add_h(doc, "4. Task 2 — VAPI System Prompt")
    add_p(doc, "Full prompt in docs/VAPI_SYSTEM_PROMPT.md. Summary of structure:")
    add_table(
        doc,
        ["Section", "Purpose"],
        [
            ["Identity & tone", "Aria, empathetic Tiger specialist, one question at a time"],
            ["System awareness", "Stage + KYC + activation variables injected at call start"],
            ["Product KB", "₹499 fee, Jewels 5=₹1, cashback tiers, KYC windows, welcome offer"],
            ["Stage routing", "8 stages, single objective per call, no skipping"],
            ["Objection paths", "7 data-backed logic branches"],
            ["Error/escalation", "Missing data, API fail, repeat objection, human request"],
            ["Write-back", "disposition, next_action, objection_codes, compliance_flags"],
            ["Hard stops", "No CVV/PIN, no verbal ID, no policy improvisation"],
        ],
    )
    add_p(doc, "VAPI variables: customer_name, current_stage, ekyc_status, vkyc_status, activation_status, retry_count, objection_history, consent_state, and 10+ more (see VAPI_SYSTEM_PROMPT.md).")

    add_h(doc, "5. Evaluation Design")
    add_p(doc, "Two layers: (A) LLM-as-judge per-call rubric; (B) operational dashboard metrics.")
    add_h(doc, "5.1 Eval rubric (6 dimensions, 0–5)", 2)
    add_bullets(
        doc,
        [
            "Stage alignment — correct objective for current_stage",
            "Data discipline — no hallucination, no prohibited disclosures",
            "Objection handling — data-backed path when objection raised",
            "Compliance — guardrails, consent, audit language",
            "Containment decision — escalate at correct threshold",
            "Customer experience — empathy, pace, clarity",
        ],
    )
    add_p(doc, "Full eval prompt: docs/EVAL_PROMPT.md")

    add_h(doc, "5.2 Operational metrics", 2)
    add_table(
        doc,
        ["Metric", "Definition", "Why"],
        [
            ["Stage completion rate", "% completing stage after agent touch", "Funnel movement"],
            ["Containment rate", "% calls without human handoff", "Cost + autonomy"],
            ["Escalation rate", "% to Inside Sales", "Agent calibration"],
            ["Avg time to activation", "Days approval → ACTIVE", "Business north star"],
            ["Drop-off recovery rate", "% stalled users re-engaged", "Core use case value"],
            ["CSAT", "Post-call satisfaction 1–5", "Experience quality"],
        ],
    )

    add_h(doc, "6. Objection support matrix")
    add_table(
        doc,
        ["Objection", "Data needed", "System behavior"],
        [
            ["Joining fee", "fee policy, welcome offer, compliance script", "Net-benefit framing; escalate after 2 exchanges"],
            ["Jewels cashback", "conversion rate, category rates", "5 Jewels=₹1 with spend example"],
            ["Low credit limit", "approved_limit, upgrade policy", "Starter limit positioning; no underwriting reveal"],
            ["Already have card", "segment, spend categories", "Use-case fit vs replacement"],
            ["Deactivation", "usage history, inactivity policy", "Minimal use-case or respectful close"],
            ["KYC complexity", "failure_reason, human assist flag", "Specific next step; human VKYC after 2 fails"],
            ["Ad mismatch", "campaign_source, offer text, disclaimer", "Clarify or escalate to compliance"],
        ],
    )

    add_h(doc, "7. Interactive prototype")
    add_p(doc, "Reviewer-facing SPA demonstrates Task 1 live: stage selector, animated system diagram, agent panel (logic/data/guardrails), objection chips, failure mode, simulated transcripts with thinking turns, eval metrics bar.")

    add_h(doc, "8. Closing")
    add_p(
        doc,
        "The voice agent is a stateful enterprise worker: it knows where the customer is, what data to trust, "
        "what action is allowed, what to write back, and when to stop. That is the core deliverable Blue Machines requested.",
    )
    return doc


def build_internal_notes() -> Document:
    doc = Document()
    style_doc(doc)
    add_title(
        doc,
        "Blue Machines — Internal Working Notes",
        "Interview prep · concepts · metrics · Rushil Kaul",
    )

    add_h(doc, "1. What they are actually testing")
    add_p(
        doc,
        "Not: build STT, telephony, or integration code. Yes: data flow, stage machine, prompt control layer, eval discipline. "
        "Think FDE = forward-deployed product operator who maps business journey → system reads/writes → measurable outcomes.",
    )

    add_h(doc, "2. The seven disciplines — definitions & interview use")
    disciplines = [
        ("Consulting", "Discovery with stakeholders, decode workflows, translate business pain to system design. Say: 'I ran discovery on drop-off points and mapped each to a READ/WRITE path.'"),
        ("Product Ops", "Own the onboarding state machine, weekly iteration on prompts/metrics, ship improvements to production. Say: 'We iterate containment and stage completion weekly from eval clusters.'"),
        ("Customer Management", "Empathy, objection handling, escalation etiquette, callback SLAs. Say: 'Customer never left without next step — disposition always written.'"),
        ("Prompt Design", "Control layer using variables + KB + stage routing — not persona fluff. Say: 'Prompt is a state router; English is the programming language.'"),
        ("AI Evals", "LLM-as-judge rubric + operational metrics; pass/fail gates on compliance. Say: 'Eval scores data_discipline and stage_alignment before CSAT.'"),
        ("Analytics", "Dashboards: containment, escalation, time-to-activation, objection breakdown. Say: 'Metrics tie to business outcome, not word quality alone.'"),
        ("Systems Integration & Workflows", "10 systems, 4 arrow types, event-driven transitions, single SoT for stage. Say: 'Agent never infers stage from chat if API has truth.'"),
    ]
    for name, desc in disciplines:
        add_h(doc, name, 2)
        add_p(doc, desc)

    add_h(doc, "3. Glossary — concepts to know cold")
    glossary = [
        ("FDE / Forward Deployed Engineer", "Embed with client, own discovery→deploy→iterate. Palantir model; BM uses for voice AI."),
        ("Stateful agent", "Behavior driven by journey stage variable, not free chat."),
        ("Disposition", "Call outcome label written to CRM: completed, callback, escalated, refused."),
        ("Containment", "Call resolved without human — key cost/scale metric."),
        ("Escalation / warm transfer", "Hand to Inside Sales with context bundle — not cold transfer."),
        ("Source of truth (SoT)", "CRM/API for stage — agent reads, never guesses."),
        ("READ/WRITE/NOTIFY/ESCALATE", "Four integration verbs in your diagram."),
        ("Guardrails", "Hard stops: no CVV, no verbal ID, no policy improvisation."),
        ("Hallucination (ops sense)", "Agent states fee/offer/status not in KB or variables."),
        ("Eval loop", "Transcript → judge → failure_reasons → prompt patch → redeploy."),
        ("Drop-off recovery", "Core Tiger use case — stalled mid-funnel users."),
        ("Jewels", "Tiger rewards currency; 5 Jewels = ₹1 — always do the math."),
        ("VKYC window", "9 AM–9 PM only — scheduling guardrail."),
        ("Context bundle", "Escalation payload: stage, objections, transcript, reason code."),
        ("Council of LLMs", "BM platform pattern — multiple models with grounding (you design prompts, not the council)."),
        ("Hot-swap DR", "Vendor failure → backup path; your failure mode simulates this."),
    ]
    add_table(doc, ["Term", "Definition"], [[a, b] for a, b in glossary])

    add_h(doc, "4. Metrics deep dive — what, why, how to use in Tiger task")
    metrics = [
        ("Stage completion rate", "% who complete current stage within 7d of agent contact", "Proves agent moves funnel", "Primary KPI per stage; compare before/after prompt change"),
        ("Containment rate", "% calls without human handoff", "Unit economics of voice AI", "Target 70–85% onboarding; drop in failure mode is expected"),
        ("Escalation rate", "% routed to Inside Sales", "Calibration signal", "Too high = weak prompt; too low = stuck customers"),
        ("Avg time to activation", "Median days approval → ACTIVE", "Business north star", "Tie agent success to revenue-ready card"),
        ("Drop-off recovery rate", "% stalled users who complete next step", "Validates drop-off use case", "Hero metric for this assignment"),
        ("CSAT", "1–5 post-call satisfaction", "Experience quality", "Lagging indicator; pair with containment"),
        ("Objection resolution rate", "% objection calls where customer continues", "Prompt + data path quality", "Slice by objection type"),
        ("FCR (first-call resolution)", "Objective achieved on first contact", "Efficiency", "Higher at APPROVED/EKYC than VKYC"),
        ("No-show rate (VKYC)", "% missed booked slots", "Scheduling effectiveness", "Drives VKYC_PENDING prompt tweaks"),
        ("Hallucination flag rate", "% calls with eval data_discipline < 3", "Compliance risk", "Must trend to zero"),
        ("Consent violation rate", "Calls without verified consent", "Regulatory hard gate", "Must be 0%"),
    ]
    add_table(doc, ["Metric", "Definition", "Why", "Tiger task use"], metrics)

    add_h(doc, "5. Blue Machines marketed flows → Tiger mapping")
    add_table(
        doc,
        ["BM case study", "Tiger parallel"],
        [
            ["MF Reactivation (SIP drop-off)", "EKYC/VKYC/Activation drop-off recovery"],
            ["Real-time data fetch", "READ Card Core, eKYC, VKYC at call start"],
            ["Objection handling", "7 objection chips with dataNeeded"],
            ["WhatsApp send", "NOTIFY arrows"],
            ["Call transfer", "ESCALATE → Inside Sales"],
            ["Cart Recovery (context memory)", "objection_history, prior_calls from CRM"],
            ["Payment Collections (empathy)", "Joining fee objection path"],
            ["Compliance (SEBI etc.)", "Compliance system + guardrails tab"],
        ],
    )

    add_h(doc, "6. Panel demo script (3 minutes)")
    add_bullets(
        doc,
        [
            "Open APPROVED — show READ arrows from Card Core + CRM",
            "Switch EKYC_PENDING — show drop-off recovery data fields",
            "Click Joining Fee objection — show Card Core fields enabling response",
            "Toggle Failure Mode — CRM offline → callback, no guessing",
            "Open Transcript — show thinking turns (agent harness)",
            "Point to Analytics bar — containment, escalation, CSAT",
            "Close: 'I also have full VAPI prompt and eval rubric in submission docs'",
        ],
    )

    add_h(doc, "7. What NOT to say")
    add_bullets(
        doc,
        [
            "I built an STT orchestrator / telephony stack",
            "The LLM decides the stage from conversation",
            "I improvised the fee policy in the prompt",
            "Metrics are just CSAT — word quality is enough",
            "Deep dive into pnpm/Vite/Netlify unless asked",
        ],
    )

    add_h(doc, "8. Likely interview questions + answer frames")
    qa = [
        ("Why not let the LLM infer stage from chat?", "CRM is SoT; inference causes redundant questions and compliance risk."),
        ("When do you escalate?", "retry>3, vkyc no_show>2, same objection twice, compliance flag, human request."),
        ("How do you prevent hallucination?", "Variables + approved KB + eval data_discipline dimension + hard stops in prompt."),
        ("How do you iterate weekly?", "Cluster eval failure_reasons → one prompt or workflow change → measure containment + stage completion."),
        ("Difference from chatbot?", "Stateful worker with WRITE-back and NOTIFY — not stateless Q&A."),
        ("How handle joining fee objection?", "READ fee policy + welcome offer; net-benefit math; escalate after 2 exchanges."),
    ]
    for q, a in qa:
        add_p(doc, f"Q: {q}", bold=True)
        add_p(doc, f"A: {a}")
        doc.add_paragraph()

    add_h(doc, "9. Assignment checklist")
    add_table(
        doc,
        ["Requirement", "Status", "Artifact"],
        [
            ["Task 1: System interaction", "Done", "Prototype + Section 3 of submission"],
            ["Task 1: READ/WRITE/NOTIFY", "Done", "SystemFlowDiagram + model.ts"],
            ["Task 1: Failure handling", "Done", "Failure Mode toggle"],
            ["Task 1: Compliance", "Done", "Guardrails tab + Compliance system"],
            ["Task 2: VAPI prompt", "Done", "docs/VAPI_SYSTEM_PROMPT.md"],
            ["Eval prompt", "Done", "docs/EVAL_PROMPT.md"],
            ["7 objections", "Done", "OBJECTIONS in model.ts"],
            ["Avoid STT/engineering", "Done", "Design-only scope"],
        ],
    )

    add_h(doc, "10. File map")
    add_bullets(
        doc,
        [
            "model.ts — all stages, systems, objections, metrics",
            "transcript.ts — simulated calls with thinking turns",
            "VAPI_SYSTEM_PROMPT.md — paste into VAPI",
            "EVAL_PROMPT.md — post-call judge",
            "GitHub: TheRuKa7/bluemachines_sample",
        ],
    )
    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    case = build_case_study()
    notes = build_internal_notes()

    paths = [
        OUT_DIR / "Blue_Machines_Case_Study_Submission_Rushil_Kaul.docx",
        OUT_DIR / "Blue_Machines_Internal_Working_Notes_Rushil_Kaul.docx",
        DOWNLOADS / "Blue_Machines_Case_Study_Submission_Rushil_Kaul.docx",
        DOWNLOADS / "Blue_Machines_Internal_Working_Notes_Rushil_Kaul.docx",
    ]
    case.save(paths[0])
    notes.save(paths[1])
    case.save(paths[2])
    notes.save(paths[3])
    print("Wrote:")
    for p in paths:
        print(f"  {p}")


if __name__ == "__main__":
    main()
